# -*- coding: utf-8 -*-
"""生成《AI预打分方案.docx》到 D:\\wqgao\\code\\AI_scoring_2。

方案要点（已与用户确认）：
- 打分主体 = 传统 CV 几何计算（零 token），AI 多模态仅作可选的可疑样本复核
- 只用分割图：无背景、无网格线，每笔一色且颜色顺序固定 → 逐笔对齐比对
- 三锚点（满分/较差/最差分割图）每维独立刻度插值映射六维分
- 不写扣分原因，输出六维分 json，J 列公式自动求和
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).parent / "AI预打分方案.docx"

DOC = Document()

style = DOC.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def heading(text: str, level: int = 1) -> None:
    h = DOC.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)


def para(text: str, bold: bool = False, indent: float = 0.0) -> None:
    p = DOC.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    return p


def bullet(text: str) -> None:
    p = DOC.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def code(text: str) -> None:
    p = DOC.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x50, 0x7F)
    p.paragraph_format.left_indent = Pt(12)


def table(headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    t = DOC.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for run in cells[i].paragraphs[0].runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    DOC.add_paragraph()


# ============ 封面标题 ============
title = DOC.add_heading("AI 预打分方案", level=0)
for run in title.runs:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
sub = DOC.add_paragraph()
run = sub.add_run(
    "基于分割图几何特征（逐笔颜色比对）的批量预打分方案 —— 为评分老师提供可修改的初稿，降低人工评分工作量。"
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ============ 一、方案定位 ============
heading("一、方案定位", 1)
para("目标：AI 做评分老师的“草稿员”。对每个新字样本批量输出六维预打分，老师只需复核和修改，而不是从零评分。", bold=False)
bullet("打分主体 = 传统 CV 几何计算（OpenCV / scikit-image），零 token 消耗，数百样本秒级完成。")
bullet("AI 多模态视觉仅作可选辅助：标记可疑样本（约 10%），不参与打分。")
bullet("预打分定位是“大概合理”的初稿，不追求最精确；最终分数以老师修改后为准。")
bullet("输出不带扣分原因，直接给六维分数，总分由 Excel 公式自动求和。")

# ============ 二、数据基础 ============
heading("二、数据基础：分割图的两个关键性质", 1)
para("所有计算只使用分割图（不使用样本图）。分割图具有两个对几何计算极为有利的性质：")
table(
    ["性质", "对评分的好处"],
    [
        ["无背景、无米字格线", "字形即前景，无需去网格线、无需去噪，预处理大幅简化"],
        ["每一笔画一个固定颜色，且颜色顺序固定", "笔画自带“身份标签”，可逐笔对齐比对，笔画规范、笔画衔接可从整字统计升级为逐笔精确计算"],
    ],
)
para("说明：笔画边缘可能存在抗锯齿过渡色，需先做颜色量化（HSV 色相聚类 / 颜色距离阈值）再分离主色，不硬编码 RGB。")

# ============ 三、总体流程 ============
heading("三、总体流程", 1)
code("锚点分割图(满分/较差/最差) ──┐")
code("                              ├─→ 颜色量化 → 逐笔分离 → 逐笔骨架化")
code("待评样本分割图 ─────────────┘")
code("                                  ↓")
code("        六维特征提取（逐笔比对 + 画布参考）")
code("                                  ↓")
code("      三锚点刻度插值映射 → 六维分（0~15/20/15/10/10/10）")
code("                                  ↓")
code("      scores_*.json → 写回 D:I → J 列公式自动求和 → K 列自动等级")
para("锚点与样本走同一条预处理与特征提取管道，保证可比性。")

# ============ 四、锚点设置 ============
heading("四、锚点设置", 1)
para("每个新字需要用户提供三张分割图作为刻度参考（同字同笔顺）：")
table(
    ["锚点", "作用", "刻度分（以该维度满分为基准）"],
    [
        ["标准满分图", "锁定该维度上限", "该维满分 Max（15/20/15/10/10/10）"],
        ["较差图", "锁定中下段尺度", "该维 40%~45%（可调）"],
        ["最差图", "锁定下限（不要求书法专业，明显差即可）", "该维 10%~15%（可调）"],
    ],
)
bullet("三个锚点可以按维度独立指定：笔画维度的“较差图”与结构维度的“较差图”可以是不同样本。")
bullet("锚点越有代表性，刻度越准；初始三张跑通后，可随评分反馈逐步补充每个维度的锚点样本。")

# ============ 五、预处理与笔画分离 ============
heading("五、预处理与笔画分离", 1)
para("分割图 RGB → 逐笔掩膜 → 逐笔骨架化，共四步：")
table(
    ["步骤", "做法", "产出"],
    [
        ["1. 颜色量化", "HSV 色相聚类或颜色距离阈值，把抗锯齿过渡色归并到主色", "主色列表 + 每像素颜色标签"],
        ["2. 逐笔分离", "按颜色标签取连通域，得到每个笔画的掩膜", "N 张单笔掩膜（N = 笔画数）"],
        ["3. 笔画数校验", "颜色数与预期笔画数对比；不一致（缺笔/粘连/异常）→ 标记复核", "异常样本名单"],
        ["4. 逐笔骨架化", "对每张单笔掩膜做细化，得到骨架与端点/交叉点", "逐笔骨架 + 拓扑点"],
    ],
)
para("异常处理：颜色数少于预期（缺笔或两笔同色粘连）时，该样本不参与插值打分，直接标记“建议老师复核”，与人工评分的 0 分规则（无效样本判 N）衔接。")

# ============ 六、六维特征设计 ============
heading("六、六维特征设计（全部基于分割图）", 1)
table(
    ["维度", "代理量（可计算）", "颜色信息利用"],
    [
        ["笔画规范 /15", "逐笔主方向角、长度、宽度均值/方差、曲率，与锚点对应笔偏差加权汇总；断笔检测", "逐笔分离"],
        ["结构规范 /20", "四宫格/八宫格墨迹占比分布、整体重心、笔间相对位置（笔 i 质心 vs 笔 j 质心）与锚点偏差", "逐笔质心"],
        ["位置规范 /15", "字形外接框中心 vs 画布中心偏移、上下左右边距比（画布为参考系）", "不涉及"],
        ["占格大小 /10", "外接框面积 ÷ 画布面积、宽高比与画布宽高比偏差", "不涉及"],
        ["笔画衔接位置 /10", "异色笔画端点两两最小间距（该接未接的缺口）、笔画交点位置 vs 锚点偏差", "核心收益"],
        ["留白空间 /10", "笔画间空隙大小分布、墨迹密度分布熵、外接框内空白方差", "笔画掩膜"],
    ],
)
para("位置规范与占格大小说明：分割图无田字格参考，以画布为参考系。预打分是粗粒度，该精度足够；老师终评时结合样本图复核。")

# ============ 七、比较与分数映射 ============
heading("七、比较与分数映射", 1)
para("每个维度独立执行三锚点刻度插值：")
code("样本特征值 f（多维代理量取加权距离）在三个锚点特征值间定位：")
code("  最差锚点特征 → 低分 L（该维满分的 10%~15%）")
code("  较差锚点特征 → 中分 M（该维满分的 40%~45%）")
code("  满分锚点特征 → 满分 Max")
code("按 f 的相对位置分段线性插值；越界截断：")
code("  f 优于满分锚点 → Max；f 劣于最差锚点 → 1 分")
code("  0 分仅用于无效样本（笔画数异常等），不参与插值")
bullet("锚点、插值斜率、低分/中分比例全部做成可调参数，便于用已有标签校准。")

# ============ 八、输出与写回 ============
heading("八、输出与写回", 1)
para("输出格式与现有流水线一致，不写扣分原因：")
code('[{"row":2,"D":9,"E":12,"F":11,"G":7,"H":6,"I":7}, ...]')
para("写回复用现有 apply_dao_scores.py：六维分写入 D:I 列，J 列 =SUM(D:I) 自动求总分，K 列公式自动出 10 档等级，L 列留给老师填最终人工评价等级，源文件不被覆盖。")

# ============ 九、可选 AI 复核 ============
heading("九、可选 AI 复核（不参与打分）", 1)
bullet("几何打分完成后，可另用 AI 视觉只标记可疑样本：笔画数异常、特征离群、分割异常。")
bullet("产出“建议老师重点复核”名单（约 10%），不产生分数、不消耗大量 token。")

# ============ 十、局限与对策 ============
heading("十、局限与对策", 1)
table(
    ["风险", "对策"],
    [
        ["抗锯齿过渡色 / 两笔同色被误分", "HSV 色相聚类 + 同色不相邻合并；笔画数异常直接标记复核"],
        ["位置、占格无田字格参考", "以画布为参考系，满足预打分粗粒度精度；老师终评看样本图"],
        ["跨字颜色方案不同", "锚点按字提供；特征全部归一化（比例、相对量）"],
        ["缺笔、粘连、分割异常", "笔画数校验拦截，标记复核，不硬算分"],
        ["特征与语义存在差距", "用锚点差值法而非绝对规则；与已有标签算 Pearson/MAE 迭代校准"],
    ],
)

# ============ 十一、开发计划 ============
heading("十一、开发计划", 1)
table(
    ["阶段", "内容", "交付物"],
    [
        ["Phase 1", "颜色量化与笔画分离库；用刀字 689 张分割图验证分离质量（颜色数与实际笔画数一致性抽查）", "stroke_separate.py + 验证报告"],
        ["Phase 2", "逐笔特征提取 + 锚点刻度插值映射", "features.py + score_mapper.py + 六维分 json"],
        ["Phase 3", "与已有 scores_new 标签评估（维度 MAE / 总分 Pearson / 等级一致率），调权重与斜率", "评估报告 + 参数配置"],
        ["Phase 4", "可选 AI 复核名单 + 对接 apply_dao_scores.py 闭环", "复核名单 + 全链路脚本"],
    ],
)
para("代码全部存放于 D:\\wqgao\\code\\AI_scoring_2，遵循既有脚本规范：显式 CLI 传参、绝对路径、不覆盖源文件、输出可配置。")

DOC.save(OUT)
print(f"Saved: {OUT}")
