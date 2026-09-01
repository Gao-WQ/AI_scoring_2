# -*- coding: utf-8 -*-
"""生成《AI预打分方案_v2.docx》到 D:\\wqgao\\code\\AI_scoring_2。

v2 相对 v1 的变更（用户要求）：
1. 新增「锚点文件格式规范」：固定目录模板 + 命名 + anchor.json，便于提供与程序读取计算
2. 需求升级为 300 字全量打分：新增「项目目录结构与批量管道」，每字输入/特征/评分/输出分开保存
3. 相应更新流程、输出写回、局限对策与开发计划
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).parent / "AI预打分方案_v2.docx"

DOC = Document()

style = DOC.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def heading(text: str, level: int = 1) -> None:
    h = DOC.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)


def para(text: str, bold: bool = False, indent: float = 0.0) -> None:
    p = DOC.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if indent:
        p.paragraph_format.left_indent = Pt(indent)
    return p


def bullet(text: str) -> None:
    p = DOC.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def code(text: str) -> None:
    p = DOC.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x50, 0x7F)
    p.paragraph_format.left_indent = Pt(12)


def table(headers: list[str], rows: list[list[str]]) -> None:
    t = DOC.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for run in cells[i].paragraphs[0].runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    DOC.add_paragraph()


# ============ 封面标题 ============
title = DOC.add_heading("AI 预打分方案（版本 2）", level=0)
for run in title.runs:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
sub = DOC.add_paragraph()
run = sub.add_run(
    "基于分割图几何特征（逐笔颜色比对）的批量预打分方案。v2 新增：锚点文件格式规范、300 字全量批量的目录与保存规范。"
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ============ 一、方案定位 ============
heading("一、方案定位", 1)
para("目标：覆盖 300 个汉字的全部样本，自动输出六维预打分，为评分老师提供可修改的初稿，降低人工评分工作量。")
bullet("打分主体 = 传统 CV 几何计算（OpenCV / scikit-image），零 token 消耗。")
bullet("AI 多模态视觉仅作可选辅助：标记可疑样本（约 10%），不参与打分。")
bullet("预打分是“大概合理”的初稿，最终分数以老师修改后为准。")
bullet("输出不带扣分原因，直接给六维分数，总分由 Excel 公式自动求和。")

# ============ 二、数据基础 ============
heading("二、数据基础：分割图的两个关键性质", 1)
para("所有计算只使用分割图（不使用样本图）。分割图具有两个对几何计算极为有利的性质：")
table(
    ["性质", "对评分的好处"],
    [
        ["无背景、无米字格线", "字形即前景，无需去网格线、无需去噪，预处理大幅简化"],
        ["每一笔画一个固定颜色，且颜色顺序固定", "笔画自带“身份标签”，可逐笔对齐比对，笔画规范、笔画衔接可从整字统计升级为逐笔精确计算"],
    ],
)
para("说明：笔画边缘可能存在抗锯齿过渡色，需先做颜色量化（HSV 色相聚类 / 颜色距离阈值）再分离主色，不硬编码 RGB。")

# ============ 三、锚点文件格式规范 ============
heading("三、锚点文件格式规范（v2 新增）", 1)
para("每个字一个锚点目录，目录名即字名。用户只需按固定模板放置三张分割图，刻度配置由程序生成默认值、按需修改。")
heading("3.1 目录与命名模板", 2)
code("AI_scoring_2\\data\\anchors\\{字}\\")
code("├── perfect.png      # 标准满分分割图（必填）")
code("├── fair.png         # 较差分割图（必填）")
code("├── worst.png        # 最差分割图（必填，明显差即可，不要求书法专业）")
code("└── anchor.json      # 刻度配置（由 init_anchor.py 自动生成，可按需修改）")
heading("3.2 图片格式要求", 2)
table(
    ["项目", "要求"],
    [
        ["文件格式", "PNG（无损，保留笔画颜色）"],
        ["色彩模式", "RGB，无透明通道混叠"],
        ["内容", "与样本同规格的分割图：无背景、每笔一色、颜色顺序固定"],
        ["尺寸", "建议与样本分割图一致；不一致时程序统一缩放归一化"],
        ["命名", "固定为 perfect.png / fair.png / worst.png，不允许别名"],
    ],
)
heading("3.3 anchor.json 结构", 2)
code('{')
code('  "char": "刀",                      // 字名，须与目录名一致')
code('  "version": "1",')
code('  "anchors": {')
code('    "perfect": {"file": "perfect.png", "score_ratio": 1.0},   // 满分锚点 → 该维满分')
code('    "fair":    {"file": "fair.png",    "score_ratio": 0.425}, // 较差锚点 → 该维 42.5%（40~45 中值）')
code('    "worst":   {"file": "worst.png",   "score_ratio": 0.125}  // 最差锚点 → 该维 12.5%（10~15 中值）')
code('  },')
code('  "dimension_overrides": {}           // 可选：某维度指定不同锚点样本或覆盖比例')
code('  // 示例：{"D": {"fair_file": "fair_stroke.png", "fair_ratio": 0.40}}')
code('}')
heading("3.4 提供与校验流程", 2)
table(
    ["步骤", "动作"],
    [
        ["1. 建目录", "运行 init_anchor.py --char 刀，自动创建 data\\anchors\\刀\\ 并生成默认 anchor.json"],
        ["2. 放图", "把三张分割图按命名放入该目录（覆盖占位说明文件）"],
        ["3. 校验", "运行校验：文件存在、PNG 可读、颜色数 ≥ 2（非空白）；失败则拒绝该字进入评分"],
        ["4. 可选调参", "按需修改 anchor.json 的 score_ratio 或 dimension_overrides"],
    ],
)
para("后续为某维度补充更多锚点样本时，在 dimension_overrides 中增加文件引用即可，无需改动主流程。")

# ============ 四、项目目录结构与批量管道 ============
heading("四、项目目录结构与批量管道（v2 新增）", 1)
para("300 字全量打分，每个字独立处理、独立保存，互不覆盖。固定目录结构如下：")
code("AI_scoring_2\\")
code("├── src\\                      # 全部代码")
code("│   ├── config.json            # 全局配置（路径、默认刻度、字符清单模式）")
code("│   ├── init_anchor.py         # 生成锚点目录模板")
code("│   ├── stroke_separate.py     # 颜色量化 + 笔画分离")
code("│   ├── features.py            # 六维特征提取")
code("│   ├── score_mapper.py        # 锚点刻度插值映射 → 六维分")
code("│   ├── apply_scores.py        # 写回 Excel（移植 apply_dao_scores.py）")
code("│   └── run_all.py             # 300 字批量管道（主入口）")
code("├── data\\")
code("│   ├── anchors\\{字}\\          # 锚点（用户提供，见第三章）")
code("│   ├── features\\{字}_features.csv   # 特征表（中间产物，每字一份）")
code("│   ├── scores\\{字}_scores.json      # 六维分结果（每字一份）")
code("│   ├── output\\{字}_all_data_new_已评分.xlsx  # 写回成品（每字一份）")
code("│   └── chars.txt              # 可选：待评字符清单，每行一字")
code("├── logs\\run_YYYYMMDD_HHMMSS.log   # 每次运行的日志与汇总")
code("└── AI预打分方案_v2.docx")
heading("4.1 输入与输出约定", 2)
table(
    ["项目", "约定"],
    [
        ["源工作簿", "直接读取 D:\\wqgao\\工作文档\\部门文件\\ori_data\\分割图\\new\\{字}_all_data_new.xlsx，不复制（路径在 config.json 可配）"],
        ["字符清单", "默认自动扫描源目录全部 *_all_data_new.xlsx；也可用 data\\chars.txt 显式指定子集"],
        ["每字保存", "features / scores / output 均按 {字} 独立命名，同字重复运行覆盖自身，不同字互不影响"],
        ["缺锚点处理", "某字无锚点目录 → 跳过并记入日志，不中断整批；补好锚点后重跑仅处理该字"],
        ["运行日志", "logs\\run_<时间戳>.log：每字状态（完成/缺锚点/异常/样本数/总分分布）+ 末尾汇总"],
        ["汇总表", "output\\summary.csv：字、样本数、总分范围/均值、等级分布，便于整体验收"],
    ],
)
heading("4.2 批量流程", 2)
code("run_all.py")
code("  1. 解析字符清单（扫描源目录 或 读 chars.txt）")
code("  2. 逐字：读 xlsx → 取全部分割图 → 预处理/笔画分离 → 六维特征 →")
code("           载入该字锚点 → 刻度插值映射 → scores json → 写回 已评分.xlsx")
code("  3. 每字结束即写日志与中间产物；全部结束输出 summary.csv")
code("单字运行：run_all.py --char 刀（补跑/单字调试用，产物与其他字互不干扰）")

# ============ 五、总体流程 ============
heading("五、总体流程", 1)
code("锚点分割图(perfect/fair/worst) ──┐")
code("                                 ├─→ 颜色量化 → 逐笔分离 → 逐笔骨架化")
code("待评样本分割图 ─────────────────┘")
code("                                     ↓")
code("          六维特征提取（逐笔比对 + 画布参考）")
code("                                     ↓")
code("        三锚点刻度插值映射 → 六维分（0~15/20/15/10/10/10）")
code("                                     ↓")
code("    scores json → 写回 D:I → J 列公式自动求和 → K 列自动等级")
para("锚点与样本走同一条预处理与特征提取管道，保证可比性。")

# ============ 六、预处理与笔画分离 ============
heading("六、预处理与笔画分离", 1)
table(
    ["步骤", "做法", "产出"],
    [
        ["1. 颜色量化", "HSV 色相聚类或颜色距离阈值，把抗锯齿过渡色归并到主色", "主色列表 + 每像素颜色标签"],
        ["2. 逐笔分离", "按颜色标签取连通域，得到每个笔画的掩膜", "N 张单笔掩膜（N = 笔画数）"],
        ["3. 笔画数校验", "颜色数与预期笔画数对比；不一致（缺笔/粘连/异常）→ 标记复核", "异常样本名单"],
        ["4. 逐笔骨架化", "对每张单笔掩膜做细化，得到骨架与端点/交叉点", "逐笔骨架 + 拓扑点"],
    ],
)
para("异常处理：颜色数少于预期（缺笔或两笔同色粘连）时，该样本不参与插值打分，直接标记“建议老师复核”，与人工评分的 0 分规则（无效样本判 N）衔接。")

# ============ 七、六维特征设计 ============
heading("七、六维特征设计（全部基于分割图）", 1)
table(
    ["维度", "代理量（可计算）", "颜色信息利用"],
    [
        ["笔画规范 /15", "逐笔主方向角、长度、宽度均值/方差、曲率，与锚点对应笔偏差加权汇总；断笔检测", "逐笔分离"],
        ["结构规范 /20", "四宫格/八宫格墨迹占比分布、整体重心、笔间相对位置（笔 i 质心 vs 笔 j 质心）与锚点偏差", "逐笔质心"],
        ["位置规范 /15", "字形外接框中心 vs 画布中心偏移、上下左右边距比（画布为参考系）", "不涉及"],
        ["占格大小 /10", "外接框面积 ÷ 画布面积、宽高比与画布宽高比偏差", "不涉及"],
        ["笔画衔接位置 /10", "异色笔画端点两两最小间距（该接未接的缺口）、笔画交点位置 vs 锚点偏差", "核心收益"],
        ["留白空间 /10", "笔画间空隙大小分布、墨迹密度分布熵、外接框内空白方差", "笔画掩膜"],
    ],
)
para("位置规范与占格大小说明：分割图无田字格参考，以画布为参考系。预打分是粗粒度，该精度足够；老师终评时结合样本图复核。")

# ============ 八、比较与分数映射 ============
heading("八、比较与分数映射", 1)
para("每个维度独立执行三锚点刻度插值：")
code("样本特征值 f（多维代理量取加权距离）在三个锚点特征值间定位：")
code("  最差锚点 → 低分（该维满分 × worst.score_ratio，默认 12.5%）")
code("  较差锚点 → 中分（该维满分 × fair.score_ratio，默认 42.5%）")
code("  满分锚点 → 满分 Max")
code("按 f 的相对位置分段线性插值；越界截断：")
code("  f 优于满分锚点 → Max；f 劣于最差锚点 → 1 分")
code("  0 分仅用于无效样本（笔画数异常等），不参与插值")
bullet("锚点、插值斜率、score_ratio 全部来自各字 anchor.json，全局默认值在 config.json，均可调。")

# ============ 九、输出与写回 ============
heading("九、输出与写回", 1)
para("每字一个结果文件，格式与现有流水线一致，不写扣分原因：")
code("data\\scores\\{字}_scores.json")
code('[{"row":2,"D":9,"E":12,"F":11,"G":7,"H":6,"I":7}, ...]')
para("写回复用 apply_dao_scores.py 逻辑（移植为 src\\apply_scores.py）：六维分写入 D:I 列，J 列 =SUM(D:I) 自动求总分，K 列公式自动出 10 档等级，L 列留给老师填最终人工评价等级，源文件不被覆盖。成品输出到 data\\output\\{字}_all_data_new_已评分.xlsx。")

# ============ 十、可选 AI 复核 ============
heading("十、可选 AI 复核（不参与打分）", 1)
bullet("几何打分完成后，可另用 AI 视觉只标记可疑样本：笔画数异常、特征离群、分割异常。")
bullet("产出每字“建议老师重点复核”名单（约 10%），随 output 一并交付，不产生分数、不消耗大量 token。")

# ============ 十一、局限与对策 ============
heading("十一、局限与对策", 1)
table(
    ["风险", "对策"],
    [
        ["抗锯齿过渡色 / 两笔同色被误分", "HSV 色相聚类 + 同色不相邻合并；笔画数异常直接标记复核"],
        ["位置、占格无田字格参考", "以画布为参考系，满足预打分粗粒度精度；老师终评看样本图"],
        ["跨字颜色方案不同", "锚点按字提供（data\\anchors\\{字}\\），特征全部归一化（比例、相对量）"],
        ["300 字锚点准备量大", "init_anchor.py 一键生成模板；缺锚点的字先跳过、补好后单字重跑，不阻塞整批"],
        ["缺笔、粘连、分割异常", "笔画数校验拦截，标记复核，不硬算分"],
        ["特征与语义存在差距", "用锚点差值法而非绝对规则；与已有标签算 Pearson/MAE 迭代校准"],
    ],
)

# ============ 十二、开发计划 ============
heading("十二、开发计划", 1)
table(
    ["阶段", "内容", "交付物"],
    [
        ["Phase 0", "目录骨架 + config.json + init_anchor.py（锚点模板）", "可运行的目录结构与锚点模板工具"],
        ["Phase 1", "颜色量化与笔画分离库；用刀字 689 张分割图验证分离质量（颜色数与实际笔画数一致性抽查）", "stroke_separate.py + 验证报告"],
        ["Phase 2", "逐笔特征提取 + 锚点刻度插值映射", "features.py + score_mapper.py + 六维分 json"],
        ["Phase 3", "写回移植 + 批量管道 run_all.py（300 字全量）", "apply_scores.py + run_all.py + summary.csv"],
        ["Phase 4", "与已有标签评估（维度 MAE / 总分 Pearson / 等级一致率），调权重与斜率", "评估报告 + 参数配置"],
        ["Phase 5", "可选 AI 复核名单", "每字复核名单"],
    ],
)
para("代码全部存放于 D:\\wqgao\\code\\AI_scoring_2，遵循既有脚本规范：显式 CLI 传参、绝对路径、不覆盖源文件、输出可配置。")

DOC.save(OUT)
print(f"Saved: {OUT}")
