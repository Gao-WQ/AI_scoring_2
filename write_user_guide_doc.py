# -*- coding: utf-8 -*-
"""生成《使用说明.docx》：与 README.md 同内容的手写汉字预打分系统使用手册（Word 版）。"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUT = Path(__file__).parent / "使用说明.docx"

DOC = Document()
style = DOC.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def heading(text: str, level: int = 1) -> None:
    h = DOC.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x63)


def para(text: str, bold: bool = False) -> None:
    p = DOC.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def code(text: str) -> None:
    p = DOC.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x50, 0x7F)
    p.paragraph_format.left_indent = Pt(12)


def table(headers: list[str], rows: list[list[str]]) -> None:
    t = DOC.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for run in cells[i].paragraphs[0].runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    DOC.add_paragraph()


# ============ 封面 ============
title = DOC.add_heading("AI 预打分系统使用说明", level=0)
for run in title.runs:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
sub = DOC.add_paragraph()
run = sub.add_run(
    "基于分割图几何特征（逐笔颜色比对）+ 三锚点刻度插值的手写汉字六维预打分系统。"
    "为评分老师提供可修改的初稿，降低人工评分工作量。版本 v3（commit 6282a72），更新记录见 CHANGELOG.md。"
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ============ 一、评分原理 ============
heading("一、评分原理（30 秒版）", 1)
for tip in [
    "每个样本有分割图（无背景、每笔一色、颜色顺序固定）→ 颜色分离出逐笔画 → 逐笔骨架分析",
    "提取六维特征（笔画/结构/位置/占格/衔接/留白）→ 与三张锚点图（满分/较差/最差）的特征做偏差比较",
    "三锚点刻度插值映射出六维分（0~15/20/15/10/10/10）→ 写入 Excel，总分由公式自动求和",
]:
    p = para("• " + tip)
    p.paragraph_format.left_indent = Pt(12)

# ============ 二、环境 ============
heading("二、环境与依赖", 1)
table(
    ["项目", "说明"],
    [
        ["Python", "D:\\app_wqgao\\anaconda3\\envs\\poc_env\\python.exe"],
        ["依赖", "openpyxl、Pillow、numpy、scipy、opencv-python、scikit-image、python-docx（均已装）"],
        ["运行目录", "所有命令在 D:\\wqgao\\code\\AI_scoring_2\\src 下执行"],
    ],
)

# ============ 三、目录结构 ============
heading("三、目录结构", 1)
code("AI_scoring_2\\")
code("├── CHANGELOG.md              # 版本更新记录")
code("├── README.md                 # 本说明（Markdown 版）")
code("├── AI预打分方案_v*.docx       # 方案文档（设计/原理）")
code("├── write_*.py                # 各 docx 的生成脚本（保留可复跑）")
code("├── src\\")
code("│   ├── main.py               # CLI 统一入口（所有命令从这里进）")
code("│   ├── config.py / config.json   # 配置加载与全局配置")
code("│   ├── init_anchor.py        # 生成锚点目录模板")
code("│   ├── apply_scores.py       # 从 scores json 单独写回 Excel")
code("│   ├── common\\              # ★ 可复用公共库")
code("│   ├── features\\            # 特征层（分离/骨架/特征/验证）")
code("│   ├── scoring\\             # 评分层（锚点插值+校准）")
code("│   └── pipeline\\            # 管道层（单字/批量）")
code("├── data\\")
code("│   ├── anchors\\{字}\\        # 锚点（每字三张图 + anchor.json）")
code("│   ├── features\\*.csv       # 特征验证表（中间产物）")
code("│   ├── scores\\{字}_scores.json   # 六维分结果")
code("│   └── output\\              # 成品 xlsx + summary.csv")
code("└── logs\\                    # 运行日志")

# ============ 四、py 文件作用 ============
heading("四、每个 py 文件的作用", 1)
table(
    ["文件", "职责", "CLI 入口"],
    [
        ["src/main.py", "CLI 统一入口，分发 5 个子命令", "✅"],
        ["src/config.py", "读取 config.json、解析绝对路径", "—"],
        ["src/init_anchor.py", "生成锚点目录模板（三张占位 + anchor.json）", "init-anchor"],
        ["src/apply_scores.py", "从已有 scores json 写回 Excel（独立入口）", "apply-scores"],
        ["src/common/io_utils.py", "目录创建、JSON 安全读写、字符清单、源工作簿扫描", "—"],
        ["src/common/image_utils.py", "图像加载、颜色 KMeans 量化、逐笔分离、骨架化、端点/交叉点", "—"],
        ["src/common/excel_utils.py", "工作簿加载、嵌入图提取、D:I 写入、公式快照与校验", "—"],
        ["src/common/anchor_utils.py", "锚点模板创建、anchor.json 解析、锚点目录校验", "—"],
        ["src/common/logging_utils.py", "控制台 + 文件日志初始化", "—"],
        ["src/features/stroke_separate.py", "颜色量化 → 逐笔掩膜 → 笔画数校验（业务层）", "—"],
        ["src/features/skeleton.py", "单笔骨架几何：方向直方图/长度/宽度/曲率/端点", "—"],
        ["src/features/features.py", "六维特征提取（layout + 逐笔）与每维偏差距离", "—"],
        ["src/features/stroke_check.py", "笔画分离质量验证（颜色数 vs 实际笔画数）", "stroke-check"],
        ["src/features/feature_check.py", "特征提取验证，输出特征表 CSV + 分布统计", "feature-check"],
        ["src/scoring/score_mapper.py", "三锚点刻度插值 → 六维分；批量映射 + 分布校准", "—"],
        ["src/pipeline/single_char.py", "单字流水线：特征 → 评分 → scores json → 写回（可选特征 sheet）", "—"],
        ["src/pipeline/run_all.py", "字符清单遍历、逐字跑、汇总 summary.csv", "run-all"],
    ],
)

# ============ 五、命令手册 ============
heading("五、命令手册", 1)
para("所有命令统一格式：python main.py <子命令> [参数]（在 src 目录下执行）。", bold=True)

heading("5.1 生成锚点模板 init-anchor", 2)
code("python main.py init-anchor --char 刀")
table(
    ["参数", "默认", "说明"],
    [
        ["--char", "刀", "字名（目录名）"],
        ["--anchor-dir", "取 config", "锚点根目录"],
    ],
)

heading("5.2 笔画分离验证 stroke-check", 2)
code("python main.py stroke-check --char 刀 --expected-strokes 2")
table(
    ["参数", "默认", "说明"],
    [
        ["--char", "刀", "验证的字"],
        ["--expected-strokes", "2", "该字预期笔画数"],
        ["--limit", "0", "抽样数，0=全部"],
        ["--n-colors", "8", "颜色量化聚类数"],
        ["--source-dir", "取 config", "源工作簿目录"],
    ],
)

heading("5.3 特征提取验证 feature-check", 2)
code("python main.py feature-check --char 刀")
para("产出 data\\features\\{字}_features.csv（16 列）+ 分布统计，用于检查特征是否有区分度。")

heading("5.4 预打分 run-all（核心命令）", 2)
code("python main.py run-all --char 上 --save-features   # 单字")
code("python main.py run-all                              # 全量（扫描源目录）")
code("python main.py run-all --chars-file data\\chars.txt # 指定清单")
table(
    ["参数", "默认", "说明"],
    [
        ["--char", "None", "只处理单字"],
        ["--chars-file", "None", "字符清单文件"],
        ["--n-colors", "8", "颜色量化聚类数"],
        ["--save-features", "False", "在成品 xlsx 新增“ai特征值”sheet"],
        ["--source-dir / --anchors-dir", "取 config", "源/锚点目录"],
    ],
)
para("产出：data\\scores\\{字}_scores.json、data\\output\\{字}_all_data_new_已评分.xlsx、summary.csv。缺锚点的字跳过不阻塞。")

heading("5.5 单独写回 apply-scores", 2)
code("python main.py apply-scores --char 上")
para("从已有 scores json 写回 Excel（评分已算好但写回失败 / 老师改分后重写时用）。")

# ============ 六、数据流 ============
heading("六、数据流", 1)
code("源 xlsx（C 列分割图）")
code("  → extract_features（六维特征，内存）")
code("  → map_scores_batch（锚点偏差 + 分布校准 → 六维分）")
code("  → data/scores/{字}_scores.json")
code("  → 写回 data/output/{字}_all_data_new_已评分.xlsx")
code("      ├── 评分汇总 sheet（D:I 六维分 + J/K/M 公式，L 留空给老师）")
code("      └──（--save-features 时）ai特征值 sheet")
heading("ai特征值 sheet（28 列，中文）", 2)
para(
    "char_id | 笔画数 | 中心偏移 | 边距不对称 | 占格面积比 | 宽高比偏差 | 密度熵 | 字内空白比 | "
    "四宫格×4 | 笔画长度均值 | 笔画宽度均值 | 笔画曲率均值 | 偏差×6 | 六维分 | 总分"
)

# ============ 七、新字流程 ============
heading("七、新字评分完整流程", 1)
code("python main.py init-anchor --char 新字")
para("→ 放入三张分割图（perfect/fair/worst.png），删占位文件，按需改 anchor.json")
code("python main.py stroke-check --char 新字 --expected-strokes N")
code("python main.py run-all --char 新字 --save-features")
para("→ 检查 data\\output\\ 成品 + summary.csv")
para("老师收到成品后：打开 xlsx → 复核“评分汇总”sheet 的六维分 → 修改 D:I（或接受）→ L 列填最终等级。")

# ============ 八、FAQ ============
heading("八、常见问题", 1)
table(
    ["问题", "处理"],
    [
        ["写回报 PermissionError", "输出 xlsx 正被 Excel/WPS 打开，关闭后重跑或 apply-scores 补写回"],
        ["分数分布整体偏低/偏高", "检查 config.json 的 calibration.enabled（默认 true）；换更合理的锚点图"],
        ["某字被跳过", "缺锚点：data\\anchors\\{字}\\ 未放三张图；补齐后 run-all --char 字 单字重跑"],
        ["报告“笔画数异常”", "分割图缺笔/粘连，该样本标记复核，不参与评分"],
        ["需要关闭分布校准", "config.json → calibration.enabled 改 false"],
        ["想回退代码版本", "见 CHANGELOG 提交历史；本地 git checkout <commit>（需确认后操作）"],
    ],
)

# ============ 九、代码约束 ============
heading("九、代码约束（开发约定）", 1)
for tip in [
    "公共函数收敛 src/common/，业务模块只 import 不复制粘贴",
    "依赖单向：pipeline → scoring/features → common",
    "路径/参数集中 config.json，代码不硬编码",
    "CLI 一律 get_args() 且每个参数带 default",
    "git 提交/推送须用户确认后执行",
]:
    p = para("• " + tip)
    p.paragraph_format.left_indent = Pt(12)

DOC.save(OUT)
print(f"Saved: {OUT}")
